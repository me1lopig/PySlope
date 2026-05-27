import streamlit as st
import pandas as pd
import io
from pyslope import Slope, Material

# --- LIBRERÍAS PARA REPORTE ---
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 1. CONFIGURACIÓN
# ==========================================
st.set_page_config(page_title="Análisis Pyslope", layout="wide")

# Diccionario de traducción
MAPA_COLORES = {
    "Marrón claro": "tan", "Amarillo arena": "khaki", "Marrón tierra": "peru",
    "Marrón oscuro": "saddlebrown", "Gris claro": "lightgrey", "Gris oscuro": "darkgrey",
    "Lila arcilla": "plum", "Verde claro": "lightgreen", "Azul claro": "lightblue",
    "Rojizo": "salmon", "Amarillo oro": "gold", "Púrpura": "purple"
}

# ==========================================
# FUNCION GENERADORA DE INFORME (DOCX)
# ==========================================
def generar_reporte(height, angle, nf, tabla_suelos, fos, fig_plotly):
    # Crear documento
    doc = Document()
    
    # --- TÍTULO ---
    titulo = doc.add_heading('Informe de Estabilidad de Taludes', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Fecha de generación: {pd.Timestamp.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph('Método de Cálculo: Bishop Simplificado')
    
    # --- 1. GEOMETRÍA Y AGUA ---
    doc.add_heading('1. Datos Geométricos', level=1)
    p = doc.add_paragraph()
    p.add_run(f'• Altura del Talud: ').bold = True
    p.add_run(f'{height} m\n')
    p.add_run(f'• Ángulo de Inclinación: ').bold = True
    p.add_run(f'{angle}°\n')
    p.add_run(f'• Nivel Freático (Profundidad): ').bold = True
    p.add_run(f'{nf} m')

    # --- 2. ESTRATIGRAFÍA ---
    doc.add_heading('2. Propiedades del Suelo', level=1)
    
    # Crear tabla en Word
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Encabezados
    hdr_cells = table.rows[0].cells
    encabezados = ['Material', 'Peso Esp. (kN/m³)', 'Cohesión (kPa)', 'Fricción (°)', 'Potencia (m)']
    for i, texto in enumerate(encabezados):
        run = hdr_cells[i].paragraphs[0].add_run(texto)
        run.bold = True
        
    # Llenar datos
    for index, row in tabla_suelos.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Material'])
        row_cells[1].text = f"{row['γ']:.2f}"
        row_cells[2].text = f"{row['c']:.2f}"
        row_cells[3].text = f"{row['φ']:.2f}"
        row_cells[4].text = f"{row['Potencia']:.2f}"

    # --- 3. RESULTADOS ---
    doc.add_heading('3. Resultados del Análisis', level=1)
    
    # Párrafo del FOS
    p_res = doc.add_paragraph()
    run_res = p_res.add_run(f'Factor de Seguridad Mínimo (FS): {fos:.3f}')
    run_res.bold = True
    run_res.font.size = Pt(14)
    run_res.font.color.rgb = RGBColor(0, 0, 139) # Azul oscuro

    # Interpretación
    if fos < 1.0:
        doc.add_paragraph('Interpretación: FALLA INMINENTE (FOS < 1.0)').bold = True
    elif fos < 1.5:
        doc.add_paragraph('Interpretación: ESTABILIDAD CONDICIONAL (1.0 < FOS < 1.5)')
    else:
        doc.add_paragraph('Interpretación: TALUD ESTABLE (FOS > 1.5)')

    # --- 4. GRÁFICO ---
    doc.add_heading('4. Superficie de Falla Rotura Crítica', level=1)
    
    try:
        # Convertir gráfico Plotly a imagen PNG en memoria
        # IMPORTANTE: Esto requiere tener instalado 'kaleido'
        img_bytes = fig_plotly.to_image(format="png", width=800, height=600, scale=2)
        image_stream = io.BytesIO(img_bytes)
        
        doc.add_picture(image_stream, width=Inches(6))
        doc.add_paragraph('Figura 1: Sección transversal mostrando estratigrafía y superficie de falla crítica.', style='Caption')
    except Exception as e:
        doc.add_paragraph(f"[No se pudo generar la imagen del gráfico. Asegúrate de instalar 'kaleido'. Error: {e}]")

    # Guardar en buffer de memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.title("⛰️ Análisis de Estabilidad mediante el Método de Bishop Simplificado")

    # ==========================================
    # ENTRADAS 
    # ==========================================
    with st.sidebar:
        st.header("1. Geometría")
        height = st.number_input("Altura Talud (m)", 1.0, 50.0, 6.0, step=0.5)
        angle = st.number_input("Ángulo (°)", 10.0, 89.0, 45.0, step=1.0)
        st.divider()
        st.header("2. Nivel Freático")
        nf_prof = st.number_input("Profundidad NF (m)", 0.0, 20.0, 3.0, step=0.5)
        
        st.divider()
        st.header("3. ⚙️ Configuración del Gráfico")
        c1, c2 = st.columns(2)
        ancho_fig = c1.number_input("Ancho (px)", 300, 2000, 800, step=50)
        alto_fig = c2.number_input("Alto (px)", 300, 2000, 600, step=50)
        
        st.subheader("Límites de los Ejes")
        usar_limites = st.checkbox("Forzar límites manuales(experimental)", value=False)
        
        if usar_limites:
            c3, c4 = st.columns(2)
            x_min = c3.number_input("X Mín (m)", value=0.0, step=1.0)
            x_max = c4.number_input("X Máx (m)", value=30.0, step=1.0)
            c5, c6 = st.columns(2)
            y_min = c5.number_input("Y Mín (m)", value=0.0, step=1.0)
            y_max = c6.number_input("Y Máx (m)", value=15.0, step=1.0)
        else:
            x_min, x_max, y_min, y_max = None, None, None, None

    with st.expander("📝 Editar Capas y Colores", expanded=True):
        df_base = pd.DataFrame([
            {"Material": "Relleno", "γ": 18.0, "c": 5.0, "φ": 28.0, "Potencia": 5.0, "Color": "Rojizo"},
            {"Material": "Arcilla", "γ": 20.0, "c": 25.0, "φ": 22.0, "Potencia": 7.0, "Color": "Amarillo arena"},
        ])
        
        tabla = st.data_editor(
            df_base, num_rows="dynamic", use_container_width=True,
            column_config={
                "γ": st.column_config.NumberColumn(format="%.1f"),
                "c": st.column_config.NumberColumn(format="%.1f"),
                "φ": st.column_config.NumberColumn(format="%.1f"),
                "Potencia": st.column_config.NumberColumn("Potencia (m)", format="%.2f", min_value=0.1, help="Espesor vertical"),
                "Color": st.column_config.SelectboxColumn(options=list(MAPA_COLORES.keys()), required=True)
            }
        )

    # ==========================================
    # CÁLCULO
    # ==========================================
    if st.button("Calcular Factor de Seguridad", type="primary"):
        if tabla.empty: st.error("Tabla vacía"); return
        
        try:
            # 1. Configuración de Pyslope
            slope = Slope(height=height, angle=angle, length=None)
            mats = []
            prof_acumulada = 0.0
            
            for i, row in tabla.iterrows():
                prof_acumulada += row["Potencia"]
                color_tecnico = MAPA_COLORES.get(row["Color"], "lightgrey")
                mats.append(Material(
                    unit_weight=row["γ"], friction_angle=row["φ"], cohesion=row["c"],
                    depth_to_bottom=prof_acumulada, name=row["Material"], color=color_tecnico
                ))
            
            slope.set_materials(*mats)
            slope.set_water_table(nf_prof)
            slope.analyse_slope()
            
            fos = slope.get_min_FOS()
            fig = slope.plot_critical()

            # 2. Configuración Visual
            fig.layout.annotations = [] 
            fig.layout.shapes = []      
            
            fig.update_layout(
                title_text="Sección Transversal",
                width=ancho_fig, height=alto_fig,
                template="plotly_white", autosize=False
            )
            
            if usar_limites:
                fig.update_xaxes(range=[x_min, x_max], title="Distancia (m)")
                fig.update_yaxes(range=[y_min, y_max], title="Elevación (m)")
                fig.update_yaxes(scaleanchor="x", scaleratio=1)
            else:
                fig.update_yaxes(scaleanchor="x", scaleratio=1, title="Elevación (m)")
                fig.update_xaxes(title="Distancia (m)")

            # 3. Mostrar Resultados en Pantalla
            st.divider()
            col_res1, col_res2 = st.columns([1, 3])
            
            with col_res1:
                st.markdown(f"### FOS: :blue[{fos:.3f}]")
                st.success("Cálculo finalizado")
                
                # --- AQUÍ GENERAMOS EL REPORTE ---
                with st.spinner("Generando informe docx..."):
                    archivo_docx = generar_reporte(height, angle, nf_prof, tabla, fos, fig)
                
                # BOTÓN DE DESCARGA
                st.download_button(
                    label="📄 Descargar Informe DOCX",
                    type="primary",
                    data=archivo_docx,
                    file_name="Informe_Estabilidad.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            with col_res2:
                st.plotly_chart(fig, use_container_width=False)

        except Exception as e:
            st.error(f"Error: {e}")
            st.info("Nota: Para generar la imagen en el reporte necesitas instalar 'kaleido' (pip install kaleido).")

if __name__ == "__main__":
    main()