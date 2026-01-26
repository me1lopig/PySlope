import streamlit as st
import pandas as pd
import io
from pyslope import Slope, Material

# --- LIBRERÍAS PARA REPORTE ---
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 1. CONFIGURACIÓN
# ==========================================
st.set_page_config(page_title="Análisis Estabilidad mediante el Método de Bishop Simplificado", layout="wide")

# Diccionario de traducción de colores
MAPA_COLORES = {
    "Marrón claro": "tan", "Amarillo arena": "khaki", "Marrón tierra": "peru",
    "Marrón oscuro": "saddlebrown", "Gris claro": "lightgrey", "Gris oscuro": "darkgrey",
    "Lila arcilla": "plum", "Verde claro": "lightgreen", "Azul claro": "lightblue",
    "Rojizo": "salmon", "Amarillo oro": "gold", "Púrpura": "purple"
}

# ==========================================
# FUNCION GENERADORA DE REPORTE
# ==========================================
def generar_reporte(height, angle, nf, tabla_suelos, fos, fig_plotly):
    doc = Document()
    
    # Título
    titulo = doc.add_heading('Informe de Estabilidad de Taludes', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Fecha: {pd.Timestamp.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph('Método: Bishop Simplificado (Pyslope)')
    
    # 1. Geometría
    doc.add_heading('1. Datos de Entrada', level=1)
    p = doc.add_paragraph()
    p.add_run(f'• Altura Talud: {height} m\n').bold = True
    p.add_run(f'• Ángulo: {angle}°\n').bold = True
    p.add_run(f'• Nivel Freático: {nf} m').bold = True

    # 2. Parámetros de Cálculo
    doc.add_heading('2. Parámetros de Cálculo', level=1)
    p2 = doc.add_paragraph()
    p2.add_run("• Método: Bishop Simplificado\n")
    p2.add_run("• Dovelas (Slices): 50\n")
    p2.add_run("• Iteraciones: 2000")

    # 3. Suelos
    doc.add_heading('3. Estratigrafía', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ['Material', 'Peso (kN/m³)', 'c (kPa)', 'φ (°)', 'Potencia (m)']
    for i, h in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        
    for index, row in tabla_suelos.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Material'])
        row_cells[1].text = f"{row['γ']:.1f}"
        row_cells[2].text = f"{row['c']:.1f}"
        row_cells[3].text = f"{row['φ']:.1f}"
        row_cells[4].text = f"{row['Potencia']:.2f}"

    # 4. Resultados
    doc.add_heading('4. Resultados', level=1)
    p_res = doc.add_paragraph()
    run_res = p_res.add_run(f'Factor de Seguridad (FS): {fos:.3f}')
    run_res.bold = True
    run_res.font.size = Pt(14)
    run_res.font.color.rgb = RGBColor(220, 20, 60)


    # 5. Imagen
    doc.add_heading('5. Gráfico de Estabilidad', level=1)
    try:
        img_bytes = fig_plotly.to_image(format="png", width=800, height=600, scale=1.5)
        doc.add_picture(io.BytesIO(img_bytes), width=Inches(6))
    except Exception as e:
        doc.add_paragraph(f"[Error generando imagen: {e}]")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.title("⛰️ Análisis de Estabilidad mediante el Método de Bishop Simplificado")

    # ==========================================
    # ENTRADAS
    # ==========================================
    with st.sidebar:
        st.header("1. Geometría")
        height = st.number_input("Altura Talud (m)", 1.0, 50.0, 6.0, step=0.5)
        angle = st.number_input("Ángulo (°)", 10.0, 89.0, 45.0, step=1.0)
        st.divider()
        st.header("2. Nivel freático")
        nf_prof = st.number_input("Profundidad NF (m)", 0.0, 20.0, 3.0, step=0.5)
        
        st.divider()
        st.header("3. ⚙️ Gráfico")
        c1, c2 = st.columns(2)
        ancho_fig = c1.number_input("Ancho (px)", 300, 2000, 800, step=50)
        alto_fig = c2.number_input("Alto (px)", 300, 2000, 600, step=50)
        
        usar_limites = st.checkbox("Forzar límites manuales (cuidado con los valores)", value=False)
        if usar_limites:
            c3, c4 = st.columns(2)
            x_min = c3.number_input("X Mín", value=0.0, step=1.0)
            x_max = c4.number_input("X Máx", value=30.0, step=1.0)
            c5, c6 = st.columns(2)
            y_min = c5.number_input("Y Mín", value=0.0, step=1.0)
            y_max = c6.number_input("Y Máx", value=15.0, step=1.0)
        else:
            x_min, x_max, y_min, y_max = None, None, None, None

    # TABLA DE MATERIALES (CORREGIDA)
    with st.expander("📝 Editar datos de entrada", expanded=True):
        df_base = pd.DataFrame([
            {"Material": "Relleno", "γ": 18.0, "c": 5.0, "φ": 28.0, "Potencia": 5.0, "Color": "Rojizo"},
            {"Material": "Arcilla", "γ": 20.0, "c": 25.0, "φ": 22.0, "Potencia": 7.0, "Color": "Amarillo arena"},
        ])
        
        tabla = st.data_editor(
            df_base, num_rows="dynamic", use_container_width=True,
            column_config={
                # Títulos de la entrada de datos
                "γ": st.column_config.NumberColumn("Peso (kN/m³)", format="%.1f", width="small"),
                "c": st.column_config.NumberColumn("Cohesión (kPa)", format="%.1f", width="small"),
                "φ": st.column_config.NumberColumn("Fricción (°)", format="%.1f", width="small"),
                "Potencia": st.column_config.NumberColumn("Potencia (m)", format="%.2f", min_value=0.1),
                "Color": st.column_config.SelectboxColumn(options=list(MAPA_COLORES.keys()), required=True)
            }
        )

    # ==========================================
    # CÁLCULO
    # ==========================================
    st.markdown("---")
    st.caption("ℹ️ **Nota Técnica:** Cálculo basado en 50 dovelas y 2000 iteraciones (Estándar Pyslope).")

    if st.button("Calcular Factor de Seguridad", type="primary"):
        if tabla.empty: st.error("Tabla vacía"); return
        
        try:
            slope = Slope(height=height, angle=angle, length=None)
            slope.update_analysis_options(
                slices=50, iterations=2000, 
                tolerance=0.005, max_iterations=50
            )
            
            mats = []
            prof_acumulada = 0.0
            for i, row in tabla.iterrows():
                prof_acumulada += row["Potencia"]
                color_tecnico = MAPA_COLORES.get(row["Color"], "lightgrey")
                mats.append(Material(
                    unit_weight=row["γ"], friction_angle=row["φ"], cohesion=row["c"],
                    depth_to_bottom=prof_acumulada, name=row["Material"], color=color_tecnico
                ))
            
            slope.set_materials(*mats)
            slope.set_water_table(nf_prof)
            
            with st.spinner(f"Analizando estabilidad..."):
                slope.analyse_slope()
                fos = slope.get_min_FOS()
                fig = slope.plot_critical()

            fig.layout.annotations = [] 
            fig.layout.shapes = []      
            
            fig.update_layout(
                title_text="Sección Transversal",
                width=ancho_fig, height=alto_fig,
                template="plotly_white", autosize=False
            )
            
            if usar_limites:
                fig.update_xaxes(range=[x_min, x_max], title="Distancia (m)")
                fig.update_yaxes(range=[y_min, y_max], title="Elevación (m)")
                fig.update_yaxes(scaleanchor="x", scaleratio=1)
            else:
                fig.update_yaxes(scaleanchor="x", scaleratio=1, title="Elevación (m)")
                fig.update_xaxes(title="Distancia (m)")

            st.divider()
            col_res1, col_res2 = st.columns([1, 3])
            
            with col_res1:
                st.markdown(f"### FS: :blue[{fos:.3f}]")
                st.success("Cálculo finalizado")
                
                with st.spinner("Creando informe..."):
                    archivo_docx = generar_reporte(height, angle, nf_prof, tabla, fos, fig)
                
                st.download_button(
                    label="📄 Descargar Informe Word",
                    data=archivo_docx,
                    file_name="Informe_Estabilidad.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

            with col_res2:
                st.plotly_chart(fig, use_container_width=False)

        except Exception as e:
            st.error(f"Error: {e}")

if __name__ == "__main__":
    main()