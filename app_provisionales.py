import streamlit as st
import pandas as pd
import numpy as np
import itertools
import io
import plotly.express as px
import plotly.graph_objects as go
import matplotlib.pyplot as plt
from pyslope import Slope, Material
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# CONFIGURACIÓN E INTERFAZ DE USUARIO
# ==========================================
st.set_page_config(page_title="GeotechParam - Taludes Provisionales VFinal", layout="wide")

st.markdown("""
    <style>
    .reportview-container .main .block-container{ max-width: 95%; }
    .stButton>button { width: 100%; }
    </style>
""", unsafe_allow_html=True)

st.title("Análisis Paramétrico y Multiescenario de Taludes Provisionales")
st.markdown("---")

if 'fos_objetivo' not in st.session_state:
    st.session_state.fos_objetivo = 1.50

tab_estratos, tab_escenarios, tab_resultados, tab_graficos, tab_informe = st.tabs([
    "1. Estratigrafía", 
    "2. Configuración de Escenarios", 
    "3. Tabla de Resultados", 
    "4. Análisis Gráfico Integrado",
    "5. Generador de Informe (Word)"
])

# ==========================================
# PESTAÑA 1: ESTRATIGRAFÍA
# ==========================================
with tab_estratos:
    st.header("Propiedades Geotécnicas de los Estratos")
    if 'df_estratos' not in st.session_state:
        st.session_state.df_estratos = pd.DataFrame({
            "Estrato": ["Capa Superior", "Capa Inferior"],
            "Prof_Base (m)": [4.0, 15.0],
            "Peso_Esp (kN/m3)": [18.0, 21.0],
            "Friccion (grados)": [26.0, 34.0],
            "Cohesion (kPa)": [12.0, 30.0]
        })
    st.session_state.df_estratos = st.data_editor(st.session_state.df_estratos, num_rows="dynamic", use_container_width=True)

# ==========================================
# PESTAÑA 2: CONFIGURACIÓN DE ESCENARIOS
# ==========================================
with tab_escenarios:
    st.header("Rangos Geométricos e Hidráulicos")
    col1, col2, col3 = st.columns(3)
    with col1:
        h_min = st.number_input("H mínima (m)", value=4.0)
        h_max = st.number_input("H máxima (m)", value=12.0)
        h_step = st.number_input("Incremento ΔH (m)", value=2.0)
    with col2:
        b_min = st.number_input("Ángulo mínimo (º)", value=30.0)
        b_max = st.number_input("Ángulo máximo (º)", value=60.0)
        b_step = st.number_input("Incremento Δβ (º)", value=15.0)
    with col3:
        zw_min = st.number_input("Zw mínima (m)", value=0.0)
        zw_max = st.number_input("Zw máxima (m)", value=4.0)
        zw_step = st.number_input("Incremento ΔZw (m)", value=2.0)

    st.session_state.fos_objetivo = st.number_input("Factor de Seguridad (FoS) de Diseño:", value=1.50, step=0.05)

# ==========================================
# MOTOR DE CÁLCULO
# ==========================================
if st.sidebar.button("🚀 Ejecutar Simulación", type="primary"):
    with st.spinner("Calculando matriz analítica mediante método de Bishop..."):
        alturas = np.arange(h_min, h_max + h_step if h_min != h_max else h_min + 0.1, h_step)
        angulos = np.arange(b_min, b_max + b_step if b_min != b_max else b_min + 0.1, b_step)
        niveles = np.arange(zw_min, zw_max + zw_step if zw_min != zw_max else zw_min + 0.1, zw_step)
        
        # Guardar rangos para el informe
        st.session_state.df_rangos = pd.DataFrame({
            "Variable Analizada": ["Altura del Talud (H)", "Inclinación (β)", "Nivel Freático (Zw)"],
            "Límite Inferior (Min)": [f"{h_min} m", f"{b_min} º", f"{zw_min} m"],
            "Límite Superior (Max)": [f"{h_max} m", f"{b_max} º", f"{zw_max} m"],
            "Incremento (Paso)": [f"{h_step} m", f"{b_step} º", f"{zw_step} m"]
        })

        combinaciones = list(itertools.product(alturas, angulos, niveles))
        lista_materiales = [Material(row["Peso_Esp (kN/m3)"], row["Friccion (grados)"], row["Cohesion (kPa)"], row["Prof_Base (m)"]) for _, row in st.session_state.df_estratos.iterrows()]
            
        resultados = []
        for H, Beta, Zw in combinaciones:
            try:
                s = Slope(height=float(H), angle=float(Beta))
                if Zw > 0: s.set_water_table(float(Zw))
                s.set_materials(*lista_materiales)
                s.analyse_slope()
                fos = s.get_cr_fos()
                estado = "Estable" if fos >= st.session_state.fos_objetivo else ("Crítico" if fos >= 1.0 else "Inestable")
                resultados.append({"Altura H (m)": H, "Inclinación β (º)": Beta, "Nivel Freático Zw (m)": Zw, "FoS": round(fos, 3), "Estado": estado})
            except:
                pass
                
        st.session_state.df_resultados = pd.DataFrame(resultados)
        st.sidebar.success("Simulación finalizada.")

# ==========================================
# PESTAÑA 3: TABLA DE RESULTADOS
# ==========================================
with tab_resultados:
    st.header("Matriz Completa de Resultados")
    if 'df_resultados' in st.session_state and not st.session_state.df_resultados.empty:
        def color_estabilidad(val):
            if val is None or isinstance(val, str): return ''
            if val < 1.0: return 'background-color: #fbcbc9; color: #900c3f; font-weight: bold;'
            elif val < st.session_state.fos_objetivo: return 'background-color: #ffeaa7; color: #d63031; font-weight: bold;'
            return 'background-color: #d4edda; color: #155724;'
        st.dataframe(st.session_state.df_resultados.style.map(color_estabilidad, subset=['FoS']), use_container_width=True)

# ==========================================
# PESTAÑA 4: GRÁFICOS (ALMACENANDO BUFFERS)
# ==========================================
with tab_graficos:
    st.header("Visualización Avanzada")
    if 'df_resultados' in st.session_state and not st.session_state.df_resultados.empty:
        df_res = st.session_state.df_resultados
        fos_lim = st.session_state.fos_objetivo
        
        st.subheader("1. Sección Transversal Crítica (Matplotlib)")
        # Encontrar el peor escenario automáticamente para el informe
        peor_caso = df_res.loc[df_res['FoS'].idxmin()]
        st.write(f"**Escenario más desfavorable calculado:** H={peor_caso['Altura H (m)']}m, β={peor_caso['Inclinación β (º)']}º, Zw={peor_caso['Nivel Freático Zw (m)']}m (FoS = {peor_caso['FoS']})")
        
        s_plot = Slope(height=float(peor_caso['Altura H (m)']), angle=float(peor_caso['Inclinación β (º)']))
        if float(peor_caso['Nivel Freático Zw (m)']) > 0: s_plot.set_water_table(float(peor_caso['Nivel Freático Zw (m)']))
        mats = [Material(r["Peso_Esp (kN/m3)"], r["Friccion (grados)"], r["Cohesion (kPa)"], r["Prof_Base (m)"]) for _, r in st.session_state.df_estratos.iterrows()]
        s_plot.set_materials(*mats)
        s_plot.analyse_slope()
        fig_perfil = s_plot.plot_boundary()
        st.pyplot(fig_perfil)
        
        # Guardar en buffer de memoria para el Word
        buf_2d = io.BytesIO()
        fig_perfil.savefig(buf_2d, format='png', bbox_inches='tight', dpi=150)
        buf_2d.seek(0)
        st.session_state.img_2d = buf_2d

        st.divider()
        
        zw_sel = st.selectbox("Filtrar Curvas de Decaimiento por Nivel Freático Zw (m):", df_res["Nivel Freático Zw (m)"].unique())
        df_plano = df_res[df_res["Nivel Freático Zw (m)"] == zw_sel].dropna()
        
        st.subheader("2. Curvas de Decaimiento (Plotly)")
        fig_lineas = px.line(df_plano, x="Altura H (m)", y="FoS", color="Inclinación β (º)", markers=True, line_shape="spline", template="plotly_white")
        fig_lineas.add_hline(y=fos_lim, line_dash="dash", line_color="red")
        fig_lineas.add_hline(y=1.0, line_dash="dot", line_color="black")
        st.plotly_chart(fig_lineas, use_container_width=True)
        
        # Intentar exportar a buffer (requiere kaleido)
        try:
            buf_lineas = io.BytesIO()
            fig_lineas.write_image(buf_lineas, format="png", width=800, height=500, scale=2)
            buf_lineas.seek(0)
            st.session_state.img_lineas = buf_lineas
        except Exception as e:
            st.warning("No se pudo exportar el gráfico estático. Asegúrese de tener 'kaleido' instalado en su entorno virtual.")
            st.session_state.img_lineas = None

# ==========================================
# PESTAÑA 5: GENERADOR DOCX
# ==========================================
def add_df_to_docx(doc, df):
    t = doc.add_table(df.shape[0]+1, df.shape[1])
    t.style = 'Table Grid'
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
    doc.add_paragraph()

with tab_informe:
    st.header("Generación de Memoria Justificativa")
    st.markdown("Se compilará un archivo Word estructurado con las tablas de configuración, los extractos de resultados y las gráficas incrustadas de forma nativa.")
    
    if 'df_resultados' in st.session_state and not st.session_state.df_resultados.empty:
        if st.button("Generar y Descargar Memoria.docx", type="primary"):
            doc = Document()
            
            # Formato de título
            t = doc.add_heading('MEMORIA TÉCNICA DE CÁLCULO: ESTABILIDAD DE TALUDES PROVISIONALES', 0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # P1
            doc.add_heading('1. OBJETO Y ALCANCE', level=1)
            p = doc.add_paragraph('El presente documento resume el análisis paramétrico de sensibilidad geométrica e hidráulica efectuado para excavaciones provisionales utilizando formulaciones de equilibrio límite mediante el método simplificado de Bishop.')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # P2
            doc.add_heading('2. PARÁMETROS GEOTÉCNICOS DEL TERRENO', level=1)
            add_df_to_docx(doc, st.session_state.df_estratos)
            
            # P3
            doc.add_heading('3. RESUMEN DE LA VARIACIÓN DE LOS PARÁMETROS', level=1)
            add_df_to_docx(doc, st.session_state.df_rangos)
            
            # P4
            doc.add_heading('4. RESULTADOS DE LA ESTABILIDAD (EXTRACTO REPRESENTATIVO)', level=1)
            doc.add_paragraph('A continuación se muestra un extracto de la iteración paramétrica, filtrando escenarios limítrofes entre la seguridad y la rotura:')
            
            # Filtrar un extracto representativo para el Word (mezcla de seguros, criticos y roturas)
            df_extracto = st.session_state.df_resultados.groupby('Estado', group_keys=False).apply(lambda x: x.sample(min(len(x), 3))).sort_values(by="Altura H (m)")
            add_df_to_docx(doc, df_extracto)
            
            # P5
            total = len(st.session_state.df_resultados)
            seguros = len(st.session_state.df_resultados[st.session_state.df_resultados['Estado'] == 'Estable'])
            criticos = len(st.session_state.df_resultados[st.session_state.df_resultados['Estado'] == 'Crítico'])
            inestables = len(st.session_state.df_resultados[st.session_state.df_resultados['Estado'] == 'Inestable'])
            
            doc.add_heading('5. RESUMEN EJECUTIVO MULTIESCENARIO', level=1)
            doc.add_paragraph(f'El modelo paramétrico ha evaluado {total} combinaciones cruzando las variables descritas.')
            doc.add_paragraph(f'- Escenarios Seguros (FoS >= {st.session_state.fos_objetivo}): {seguros}', style='List Bullet')
            doc.add_paragraph(f'- Escenarios Críticos: {criticos}', style='List Bullet')
            doc.add_paragraph(f'- Escenarios Inestables (Rotura): {inestables}', style='List Bullet')
            doc.add_paragraph()
            
            # P6 - Gráficos Incrustados
            doc.add_heading('6. ANÁLISIS GRÁFICO DE SENSIBILIDAD', level=1)
            
            doc.add_heading('6.1. Sección Transversal del Escenario Crítico', level=2)
            if 'img_2d' in st.session_state and st.session_state.img_2d is not None:
                doc.add_picture(st.session_state.img_2d, width=Inches(6.0))
            
            doc.add_heading('6.2. Curvas de Decaimiento Paramétrico', level=2)
            if 'img_lineas' in st.session_state and st.session_state.img_lineas is not None:
                doc.add_picture(st.session_state.img_lineas, width=Inches(6.0))
            
            # Guardar Word en Buffer
            buf_docx = io.BytesIO()
            doc.save(buf_docx)
            buf_docx.seek(0)
            
            st.download_button(
                label="📥 Descargar Memoria.docx Completa",
                data=buf_docx,
                file_name="Memoria_Taludes_Integrada.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )